from io import BytesIO
from docx import Document
from docx.shared import Pt


def create_word(report_text: str, file_path: str) -> None:
    doc = Document()
    doc.add_heading("Referto Radiologico", level=1)
    p = doc.add_paragraph(report_text)
    for run in p.runs:
        run.font.size = Pt(11)
    disclaimer = (
        "Questo referto è generato con supporto AI e non sostituisce il giudizio medico. "
        "Evitare l'uso di dati identificativi. Valutazione clinica finale a cura del medico responsabile."
    )
    doc.add_paragraph("")
    doc.add_paragraph(disclaimer)
    doc.save(file_path)