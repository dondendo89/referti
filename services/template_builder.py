from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_tosca_template(output_path: str):
    doc = Document()
    
    # Margini
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Intestazione Prescrittore (Dr. Tosca)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Dr. med. Enrico Tosca\nVia Lugano 8\n6988 Ponte Tresa\nenrico.tosca@hin.ch")
    run.font.size = Pt(10)
    run.font.name = 'Arial'

    doc.add_paragraph("\n\n")

    # Placeholder per il contenuto
    p = doc.add_paragraph("Paziente: [PAZIENTE_NOME] [PAZIENTE_DATA_NASCITA]")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.style = 'Heading 1'

    doc.add_paragraph("\n")
    
    # Corpo del referto (sarà sostituito dinamicamente)
    doc.add_paragraph("[CONTENUTO_REFERTO]")
    
    doc.add_paragraph("\n\n")

    # Firma
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Cordiali saluti,\n\nDr. med. Riccardo La Macchia\nFMH Radiologia Medica")
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    doc.save(output_path)

def create_mantini_template(output_path: str):
    doc = Document()
    
    # Margini
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # Intestazione Prescrittore (Dr. Mantini)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Dr. med. Vito Mantini\nVia Vedeggio 1\n6814 Lamone\nvitomantini@hin.ch")
    run.font.size = Pt(10)
    run.font.name = 'Arial'

    doc.add_paragraph("\n\n")

    # Placeholder per il contenuto
    p = doc.add_paragraph("Paziente: [PAZIENTE_NOME] [PAZIENTE_DATA_NASCITA]")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.style = 'Heading 1'

    doc.add_paragraph("\n")
    
    # Corpo del referto (sarà sostituito dinamicamente)
    doc.add_paragraph("[CONTENUTO_REFERTO]")
    
    doc.add_paragraph("\n\n")

    # Firma
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Cordiali saluti,\n\nDr. med. Riccardo La Macchia\nFMH Radiologia Medica\nSpec. Mini-interventistica muscoloscheletrica")
    run.font.size = Pt(11)
    run.font.name = 'Arial'

    doc.save(output_path)
